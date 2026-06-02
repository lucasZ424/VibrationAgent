from pathlib import Path

from docx import Document

from vibration_agent.ingestion.classify import classify_document, iter_supported_files, scan_inputs


def test_corrupt_docx_classifies_as_docx_with_inspection_warning(tmp_path):
    docx = tmp_path / "manual.docx"
    docx.write_bytes(b"not a real docx, classification only")

    result = classify_document(docx)

    assert result.kind == "docx"
    assert result.processing_strategy == "docx"
    assert result.page_count == 0
    assert any("DOCX could not be inspected" in warning for warning in result.warnings)


def test_docx_classify_surfaces_unexpected_inspection_errors(tmp_path, monkeypatch):
    import vibration_agent.ingestion.docx_parser as docx_parser

    docx = tmp_path / "manual.docx"
    document = Document()
    document.add_paragraph("Rotor damping note")
    document.save(docx)

    def fail_unexpected(*args, **kwargs):
        raise RuntimeError("unexpected parser bug")

    monkeypatch.setattr(docx_parser, "inspect_docx", fail_unexpected)

    try:
        classify_document(docx)
    except RuntimeError as exc:
        assert "unexpected parser bug" in str(exc)
    else:
        raise AssertionError("unexpected DOCX inspection errors should fail loud")


def test_iter_supported_files_includes_docx_and_skips_office_lockfiles(tmp_path):
    supported = tmp_path / "note.txt"
    docx = tmp_path / "manual.docx"
    lockfile = tmp_path / "~$manual.docx"
    supported.write_text("rotor note", encoding="utf-8")
    document = Document()
    document.add_paragraph("Rotor damping note")
    document.save(docx)
    lockfile.write_bytes(b"lock")

    files = list(iter_supported_files(tmp_path))
    scanned = scan_inputs(tmp_path)

    assert files == [docx.resolve(), supported.resolve()]
    assert [item.filename for item in scanned] == ["manual.docx", "note.txt"]


def test_language_detection_for_mixed_engineering_text(tmp_path):
    text = tmp_path / "mixed.md"
    text.write_text("转子轴承振动诊断 rotor", encoding="utf-8")

    result = classify_document(text)

    assert result.kind == "text"
    assert result.processing_strategy == "text"
    assert result.language == "mixed"
