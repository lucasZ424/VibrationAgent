import base64
from pathlib import Path
from types import SimpleNamespace

from docx import Document
from docx.enum.text import WD_BREAK

import vibration_agent.ingestion.docx_parser as docx_parser
from vibration_agent.ingestion.chunking import chunk_pages
from vibration_agent.ingestion.docx_parser import DocxParseError, docx_page_count, extract_docx_text, parse_docx


_PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
)


def _write_docx(path: Path, *, with_image: bool = False) -> None:
    document = Document()
    document.add_heading("Rotor damping note", level=1)
    document.add_paragraph("Damping reduces resonant vibration near critical speed.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Signal"
    table.cell(0, 1).text = "Purpose"
    table.cell(1, 0).text = "Orbit"
    table.cell(1, 1).text = "Rotor vibration diagnosis"
    if with_image:
        image_path = path.with_suffix(".png")
        image_path.write_bytes(_PNG_1X1)
        document.add_picture(str(image_path))
    document.save(path)


def test_parse_docx_outputs_common_page_schema_with_table_and_image(tmp_path):
    docx = tmp_path / "note.docx"
    _write_docx(docx, with_image=True)

    pages = parse_docx(docx, doc_id="doc1", asset_dir=tmp_path / "assets")

    assert len(pages) == 1
    page = pages[0]
    assert page.primary_engine == "python-docx"
    assert page.layout_quality == "ok"
    assert "critical speed" in page.normalized_text
    assert page.blocks[0].block_type == "title"
    assert any(block.block_type == "table" for block in page.blocks)
    assert {asset.object_type for asset in page.assets} == {"table", "figure"}
    assert all(asset.asset_path for asset in page.assets)
    assert page.metadata["docx_pagination"]["effective"] == "logical"
    assert all("anchor" in asset.metadata for asset in page.assets)


def test_extract_docx_text_includes_tables(tmp_path):
    docx = tmp_path / "note.docx"
    _write_docx(docx)

    text = extract_docx_text(docx)

    assert "Damping reduces resonant vibration" in text
    assert "Signal | Purpose" in text


def test_docx_table_text_is_retrievable_through_chunk_text(tmp_path):
    docx = tmp_path / "note.docx"
    _write_docx(docx)

    pages = parse_docx(docx, doc_id="doc1")
    chunks = chunk_pages(
        pages,
        doc_id="doc1",
        title="Rotor damping note",
        source_path=docx,
        source_type="book",
        target_tokens=600,
        overlap_tokens=60,
    )

    assert chunks
    assert "Orbit | Rotor vibration diagnosis" in chunks[0]["text"]


def test_docx_page_count_is_one_logical_page_even_with_word_page_break(tmp_path):
    docx = tmp_path / "note.docx"
    document = Document()
    paragraph = document.add_paragraph("Before page break")
    paragraph.add_run().add_break(WD_BREAK.PAGE)
    paragraph.add_run("after page break")
    document.save(docx)

    assert docx_page_count(docx) == 1


def test_rendered_docx_pagination_falls_back_when_soffice_missing(tmp_path, monkeypatch):
    docx = tmp_path / "note.docx"
    _write_docx(docx)
    monkeypatch.setattr(docx_parser, "_soffice_path", lambda explicit_path=None: None)

    pages = parse_docx(docx, doc_id="doc1", pagination_mode="rendered")

    assert len(pages) == 1
    pagination = pages[0].metadata["docx_pagination"]
    assert pagination["requested"] == "rendered"
    assert pagination["effective"] == "logical_fallback"
    assert pagination["reason"] == "render_unavailable"
    assert pages[0].metadata["warnings"]


def test_rendered_docx_warning_redacts_local_paths_from_soffice_stderr(tmp_path, monkeypatch):
    docx = tmp_path / "note.docx"
    _write_docx(docx)
    monkeypatch.setattr(docx_parser, "_soffice_path", lambda explicit_path=None: "soffice")

    def failed_runner(*args, **kwargs):
        return SimpleNamespace(returncode=1, stderr=r"failed C:\Users\me\secret.docx /home/me/private/out.pdf")

    pages = parse_docx(docx, doc_id="doc1", pagination_mode="rendered", render_runner=failed_runner)

    warning = pages[0].metadata["warnings"][0]
    assert "[REDACTED_LOCAL_PATH]" in warning
    assert "C:\\Users\\me" not in warning
    assert "/home/me" not in warning


def test_rendered_docx_pagination_falls_back_without_layout_mapping(tmp_path, monkeypatch):
    docx = tmp_path / "note.docx"
    rendered_pdf = tmp_path / "note.pdf"
    _write_docx(docx)
    monkeypatch.setattr(docx_parser, "_pdf_page_count", lambda path: 3)

    pages = parse_docx(docx, doc_id="doc1", pagination_mode="rendered", rendered_pdf_path=rendered_pdf)

    assert len(pages) == 1
    page = pages[0]
    pagination = page.metadata["docx_pagination"]
    assert pagination["effective"] == "logical_fallback"
    assert pagination["reason"] == "missing_layout_mapping"
    assert pagination["rendered_page_count"] == 3
    assert page.assets[0].metadata["anchor"]["page_anchor_type"] == "logical_fallback"


def test_rendered_single_page_docx_records_rich_asset_anchor_metadata(tmp_path, monkeypatch):
    docx = tmp_path / "note.docx"
    rendered_pdf = tmp_path / "note.pdf"
    _write_docx(docx, with_image=True)
    monkeypatch.setattr(docx_parser, "_pdf_page_count", lambda path: 1)

    pages = parse_docx(
        docx,
        doc_id="doc1",
        asset_dir=tmp_path / "assets",
        pagination_mode="rendered",
        rendered_pdf_path=rendered_pdf,
    )

    page = pages[0]
    assert page.metadata["docx_pagination"]["effective"] == "rendered_single_page"
    anchors = [asset.metadata["anchor"] for asset in page.assets]
    assert {anchor["page_anchor_type"] for anchor in anchors} == {"rendered_single_page"}
    assert all(anchor["rendered_page_no"] == 1 for anchor in anchors)
    assert any(anchor["source"] == "docx_image_relationship" for anchor in anchors)


def test_empty_docx_raises_parse_error(tmp_path):
    docx = tmp_path / "empty.docx"
    Document().save(docx)

    try:
        parse_docx(docx, doc_id="doc1")
    except DocxParseError as exc:
        assert "no extractable" in str(exc)
    else:
        raise AssertionError("empty DOCX should be insufficient")


def test_corrupt_docx_raises_parse_error(tmp_path):
    docx = tmp_path / "corrupt.docx"
    docx.write_bytes(b"not a zip")

    try:
        parse_docx(docx, doc_id="doc1")
    except DocxParseError as exc:
        assert "could not be opened" in str(exc)
    else:
        raise AssertionError("corrupt DOCX should be insufficient")
