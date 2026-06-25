import json
from pathlib import Path

import fitz
from docx import Document

from vibration_agent.config import load
from vibration_agent.ingestion.classify import classify_document
from vibration_agent.ingestion import pipeline
from vibration_agent.ingestion.ocr import router
from vibration_agent.ingestion.pipeline import chunk_document_pages, chunk_documents, parse_document_pages
from vibration_agent.ingestion.pymupdf_parser import parse_native_pdf
from vibration_agent.schemas import IngestionManifest


def _make_text_pdf(path: Path) -> None:
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Rotor imbalance produces synchronous vibration at running speed. " * 4)
    doc.save(path)
    doc.close()


def _make_blank_pdf(path: Path) -> None:
    doc = fitz.open()
    doc.new_page()
    doc.save(path)
    doc.close()


def test_parse_native_pdf_pages(tmp_path):
    pdf = tmp_path / "native.pdf"
    _make_text_pdf(pdf)

    pages = parse_native_pdf(pdf, doc_id="doc1")

    assert len(pages) == 1
    assert pages[0].primary_engine == "pymupdf"
    assert pages[0].layout_quality == "ok"
    assert "Rotor imbalance" in pages[0].normalized_text
    assert pages[0].blocks


def test_parse_document_pages_writes_jsonl(tmp_path):
    pdf = tmp_path / "native.pdf"
    _make_text_pdf(pdf)
    doc = classify_document(pdf)
    settings = load()
    settings.paths.ocr_dir = tmp_path / "ocr"

    result = parse_document_pages(doc, settings=settings, max_pages=1, write_output=True)
    output_path = Path(result["output_path"])
    rows = [json.loads(line) for line in output_path.read_text(encoding="utf-8").splitlines()]

    assert result["status"] == "ok"
    assert output_path.exists()
    assert rows[0]["doc_id"] == doc.doc_id
    assert rows[0]["primary_engine"] == "pymupdf"


def test_bulk_native_ingestion_keeps_optional_region_ocr_disabled(tmp_path, monkeypatch):
    # WHY: retaining figures must not implicitly launch expensive PaddleOCR for every image.
    pdf = tmp_path / "native.pdf"
    _make_text_pdf(pdf)
    doc = classify_document(pdf)
    settings = load()
    settings.paths.ocr_dir = tmp_path / "ocr"
    settings.paths.extracted_dir = tmp_path / "extracted"
    captured = {}

    monkeypatch.setattr(
        pipeline,
        "parse_native_pdf",
        lambda *args, **kwargs: captured.update(kwargs) or [],
    )

    parse_document_pages(doc, settings=settings, write_output=False)

    assert settings.visual_recovery.enabled is True
    assert settings.visual_recovery.region_ocr_enabled is False
    assert captured["page_ocr_enabled"] is True
    assert captured["region_ocr_enabled"] is False


def test_ocr_router_returns_review_page_when_paddle_fails(tmp_path, monkeypatch):
    pdf = tmp_path / "scan.pdf"
    _make_blank_pdf(pdf)

    def fail_paddle(*args, **kwargs):
        raise RuntimeError("forced paddle failure")

    monkeypatch.setattr(router.paddle_engine, "run", fail_paddle)
    page = router.ocr_page(pdf, 1, doc_id="scan1", image_dir=tmp_path / "images")

    assert page.page_no == 1
    assert page.needs_review is True
    assert page.fallback_used is True
    assert page.layout_quality in {"empty", "low"}


def test_chunk_document_pages_writes_formal_chunk_outputs(tmp_path):
    pdf = tmp_path / "native.pdf"
    _make_text_pdf(pdf)
    doc = classify_document(pdf)
    settings = load()
    settings.paths.ocr_dir = tmp_path / "ocr"
    settings.paths.chunks_dir = tmp_path / "chunks"
    settings.paths.exports_dir = tmp_path / "exports"
    settings.paths.extracted_dir = tmp_path / "extracted"

    result = chunk_document_pages(doc, settings=settings, max_pages=1, write_output=True)

    assert result["stage"] == "document_structured_export"
    assert result["chunk_count"] >= 1
    assert Path(result["chunks_output_path"]).exists()
    assert Path(result["api_context_output_path"]).exists()
    assert Path(result["manifest_output_path"]).exists()
    chunk = json.loads(Path(result["chunks_output_path"]).read_text(encoding="utf-8").splitlines()[0])
    assert chunk["assets"][0]["object_type"] == "body"
    assert chunk["pages"] == [1]


def test_docx_parse_and_chunk_outputs_match_common_shape(tmp_path):
    source = tmp_path / "rotor_note.docx"
    document = Document()
    document.add_heading("Rotor damping", level=1)
    document.add_paragraph("Damping reduces resonant vibration near critical speed. " * 3)
    document.save(source)
    doc = classify_document(source)
    settings = load()
    settings.paths.ocr_dir = tmp_path / "ocr"
    settings.paths.chunks_dir = tmp_path / "chunks"
    settings.paths.exports_dir = tmp_path / "exports"
    settings.paths.extracted_dir = tmp_path / "extracted"

    page_result = parse_document_pages(doc, settings=settings, max_pages=1, write_output=True)
    chunk_result = chunk_document_pages(doc, settings=settings, max_pages=1, write_output=True)

    assert page_result["status"] == "ok"
    assert page_result["pages"][0]["primary_engine"] == "python-docx"
    assert chunk_result["status"] == "ok"
    chunk = json.loads(Path(chunk_result["chunks_output_path"]).read_text(encoding="utf-8").splitlines()[0])
    assert chunk["pages"] == [1]
    assert chunk["source_path"].endswith("rotor_note.docx")


def test_docx_parse_defaults_missing_page_count_to_one_logical_page(tmp_path):
    source = tmp_path / "rotor_note.docx"
    document = Document()
    document.add_paragraph("Damping reduces resonant vibration near critical speed. " * 3)
    document.save(source)
    doc = classify_document(source).model_copy(update={"page_count": None})
    settings = load()
    settings.paths.ocr_dir = tmp_path / "ocr"

    result = parse_document_pages(doc, settings=settings, max_pages=None, write_output=False)

    assert result["status"] == "ok"
    assert result["page_count"] == 1
    assert result["processed_pages"] == 1


def test_corrupt_docx_returns_structured_insufficient(tmp_path):
    source = tmp_path / "corrupt.docx"
    source.write_bytes(b"not a zip")
    doc = classify_document(source)
    settings = load()
    settings.paths.ocr_dir = tmp_path / "ocr"

    result = parse_document_pages(doc, settings=settings, write_output=True)

    assert result["status"] == "insufficient"
    assert result["pages"] == []
    assert result["warnings"]



def test_chunk_document_pages_writes_objective9_manifest(tmp_path):
    pdf = tmp_path / "native.pdf"
    _make_text_pdf(pdf)
    doc = classify_document(pdf)
    settings = load()
    settings.paths.ocr_dir = tmp_path / "ocr"
    settings.paths.chunks_dir = tmp_path / "chunks"
    settings.paths.exports_dir = tmp_path / "exports"
    settings.paths.extracted_dir = tmp_path / "extracted"

    result = chunk_document_pages(doc, settings=settings, max_pages=1, write_output=True)
    manifest_path = Path(result["manifest_output_path"])
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))

    assert manifest["type"] == "document_ingestion_manifest"
    assert manifest["doc_id"] == doc.doc_id
    assert manifest["input"]["source_path"] == doc.source_path
    assert manifest["input"]["sha256"] == doc.sha256
    assert manifest["input"]["language"] == doc.language
    assert manifest["input"]["doc_id_mode"] == "content"
    assert manifest["counts"]["processed_pages"] == 1
    assert manifest["counts"]["chunk_count"] == result["chunk_count"]
    assert manifest["needs_review_pages"] == []
    assert "quality" in manifest
    assert manifest["quality"]["page_ocr_confidence_min"] == 1.0
    assert manifest["markdown_outputs"] == []
    assert set(manifest["outputs"]) == {"pages_jsonl", "chunks_jsonl", "api_context_json", "manifest_json"}
    assert all(Path(path).exists() for path in manifest["outputs"].values())
    assert Path(manifest["outputs"]["pages_jsonl"]).parent == tmp_path / "ocr" / "book" / doc.doc_id
    assert not list(tmp_path.rglob("*.md"))

    page_row = json.loads(Path(manifest["outputs"]["pages_jsonl"]).read_text(encoding="utf-8").splitlines()[0])
    chunk_row = json.loads(Path(manifest["outputs"]["chunks_jsonl"]).read_text(encoding="utf-8").splitlines()[0])
    assert page_row["schema_version"] == "0.1"
    assert chunk_row["schema_version"] == "0.1"
    assert chunk_row["doc_id"] == doc.doc_id

def test_chunk_documents_batch_writes_manifests(tmp_path):
    raw_dir = tmp_path / "raw"
    raw_dir.mkdir()
    pdf = raw_dir / "native.pdf"
    _make_text_pdf(pdf)
    settings = load()
    settings.paths.ocr_dir = tmp_path / "ocr"
    settings.paths.chunks_dir = tmp_path / "chunks"
    settings.paths.exports_dir = tmp_path / "exports"
    settings.paths.extracted_dir = tmp_path / "extracted"

    result = chunk_documents(raw_dir, settings=settings, max_pages=1, write_output=True)

    assert result["stage"] == "document_structured_export_batch"
    assert result["document_count"] == 1
    manifest_path = Path(result["documents"][0]["manifest_output_path"])
    assert manifest_path.exists()
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    assert manifest["outputs"]["manifest_json"] == str(manifest_path)
    assert IngestionManifest.model_validate(manifest).counts.chunk_count == result["documents"][0]["chunk_count"]
