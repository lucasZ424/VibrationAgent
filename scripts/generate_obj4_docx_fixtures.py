"""Regenerate Phase-2 Obj4 DOCX ingestion fixtures."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))

from vibration_agent.ingestion.chunking import chunk_pages, write_jsonl  # noqa: E402
from vibration_agent.ingestion.docx_parser import parse_docx  # noqa: E402


DOC_ID = "fixture_rotor_zh_docx"
DOC_TITLE = "\u8f6c\u5b50\u963b\u5c3c\u4e0e\u4e34\u754c\u8f6c\u901f"
BODY_1 = (
    "\u963b\u5c3c\u6bd4 \u03b6 \u8868\u793a\u7cfb\u7edf\u8017\u6563\u80fd\u91cf\u7684"
    "\u80fd\u529b\u3002\u963b\u5c3c\u6bd4\u589e\u52a0\u65f6\uff0c\u81ea\u7531\u632f\u52a8"
    "\u8870\u51cf\u66f4\u5feb\u3002"
)
BODY_2 = (
    "\u8f6c\u5b50\u63a5\u8fd1\u4e34\u754c\u8f6c\u901f\u65f6\uff0c\u54cd\u5e94"
    "\u5e45\u503c\u4f1a\u653e\u5927\uff0c\u76f8\u4f4d\u4e5f\u4f1a\u5feb\u901f"
    "\u53d8\u5316\u3002"
)
TABLE_ROWS = (
    ("\u76d1\u6d4b\u91cf", "\u7528\u9014"),
    ("\u632f\u5e45\u548c\u8f74\u5fc3\u8f68\u8ff9", "\u5224\u65ad\u8f6c\u5b50\u632f\u52a8\u98ce\u9669"),
)


def write_docx(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading(DOC_TITLE, level=1)
    document.add_paragraph(BODY_1)
    document.add_paragraph(BODY_2)
    table = document.add_table(rows=len(TABLE_ROWS), cols=len(TABLE_ROWS[0]))
    for row_index, row in enumerate(TABLE_ROWS):
        for col_index, text in enumerate(row):
            table.cell(row_index, col_index).text = text
    document.save(path)


def main() -> int:
    fixtures = ROOT / "tests" / "fixtures"
    raw_docx = fixtures / "raw" / "small_vibration_zh.docx"
    write_docx(raw_docx)

    pages = parse_docx(raw_docx, doc_id=DOC_ID, extract_image_assets=False)
    page_rows = [{"schema_version": "0.1", **page.model_dump(mode="json")} for page in pages]
    write_jsonl(fixtures / "ocr" / "sample_zh_docx_pages.jsonl", page_rows)

    chunks = chunk_pages(
        pages,
        doc_id=DOC_ID,
        title=DOC_TITLE,
        source_path="tests/fixtures/raw/small_vibration_zh.docx",
        source_type="book",
        target_tokens=600,
        overlap_tokens=60,
    )
    write_jsonl(fixtures / "chunks" / "sample_zh_docx_chunks.jsonl", chunks)

    print(f"wrote {raw_docx}")
    print(f"docx_chunks={len(chunks)} pages={chunks[0]['pages'] if chunks else []}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
