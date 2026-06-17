"""DOCX parser producing the common page-level schema."""
from __future__ import annotations

import re
import shutil
import subprocess
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Callable, Iterable

from vibration_agent.ingestion.assets import asset_anchor_metadata
from vibration_agent.ingestion.layout import classify_text_block, make_asset_id
from vibration_agent.ingestion.pymupdf_parser import normalize_text
from vibration_agent.schemas import DocumentAsset, OcrPage, PageBlock


class DocxParseError(RuntimeError):
    """Raised when a DOCX cannot be parsed into usable page content."""


class DocxRenderError(RuntimeError):
    """Raised when optional DOCX rendering cannot produce a PDF page count."""


@dataclass(frozen=True)
class DocxRenderResult:
    pdf_path: Path
    page_count: int
    backend: str = "soffice"


_WINDOWS_PATH_RE = re.compile(r"[A-Za-z]:\\[^\s\"']+")
_POSIX_PATH_RE = re.compile(r"(?<![\w.])/(?:Users|home|var|tmp|mnt)/[^\s\"']+")


def _redact_local_paths(text: str) -> str:
    text = _WINDOWS_PATH_RE.sub("[REDACTED_LOCAL_PATH]", text)
    return _POSIX_PATH_RE.sub("[REDACTED_LOCAL_PATH]", text)


def _load_document(path: str | Path):
    try:
        from docx import Document  # type: ignore
    except ModuleNotFoundError as exc:
        raise DocxParseError("python-docx is required for DOCX parsing.") from exc

    try:
        return Document(str(Path(path).resolve()))
    except Exception as exc:
        raise DocxParseError(f"DOCX could not be opened: {exc}") from exc


def _style_name(paragraph) -> str:
    style = getattr(paragraph, "style", None)
    return str(getattr(style, "name", "") or "").lower()


def _paragraph_kind(text: str, paragraph, *, block_index: int) -> str:
    style = _style_name(paragraph)
    if style.startswith("heading") or style in {"title", "subtitle"}:
        return "title"
    return classify_text_block(text, block_index=block_index)


def _iter_table_texts(document) -> Iterable[str]:
    for table in document.tables:
        rows: list[str] = []
        for row in table.rows:
            cells = [normalize_text(cell.text) for cell in row.cells if normalize_text(cell.text)]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            yield "\n".join(rows)


def _text_parts_from_document(document) -> list[str]:
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = normalize_text(paragraph.text)
        if text:
            parts.append(text)
    parts.extend(_iter_table_texts(document))
    return parts


def _image_extension(content_type: str) -> str:
    subtype = content_type.rsplit("/", 1)[-1].lower()
    return {
        "jpeg": ".jpg",
        "jpg": ".jpg",
        "png": ".png",
        "gif": ".gif",
        "bmp": ".bmp",
        "tiff": ".tiff",
        "webp": ".webp",
    }.get(subtype, ".bin")


def _soffice_path(explicit_path: str | Path | None = None) -> str | None:
    if explicit_path:
        path = str(explicit_path)
        return path if Path(path).exists() else shutil.which(path)
    return shutil.which("soffice") or shutil.which("libreoffice")


def _pdf_page_count(pdf_path: str | Path) -> int:
    try:
        import fitz  # type: ignore
    except ModuleNotFoundError as exc:
        raise DocxRenderError("PyMuPDF is required to inspect rendered DOCX PDFs.") from exc

    try:
        with fitz.open(Path(pdf_path).resolve()) as pdf:
            return int(pdf.page_count)
    except Exception as exc:
        raise DocxRenderError(f"Rendered DOCX PDF could not be inspected: {exc}") from exc


def render_docx_to_pdf(
    docx_path: str | Path,
    *,
    output_dir: str | Path,
    soffice_path: str | Path | None = None,
    runner: Callable[..., Any] = subprocess.run,
) -> DocxRenderResult:
    """Render DOCX to PDF with headless LibreOffice when explicitly requested."""
    source = Path(docx_path).resolve()
    destination = Path(output_dir).resolve()
    destination.mkdir(parents=True, exist_ok=True)
    executable = _soffice_path(soffice_path)
    if executable is None:
        raise DocxRenderError("LibreOffice soffice executable not found.")

    command = [
        executable,
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(destination),
        str(source),
    ]
    try:
        result = runner(command, check=False, capture_output=True, text=True)
    except Exception as exc:
        raise DocxRenderError(f"LibreOffice soffice failed to start: {exc}") from exc
    return_code = int(getattr(result, "returncode", 0) or 0)
    if return_code != 0:
        stderr = str(getattr(result, "stderr", "") or "").strip()
        raise DocxRenderError(f"LibreOffice soffice failed with code {return_code}: {stderr}")

    pdf_path = destination / f"{source.stem}.pdf"
    if not pdf_path.exists():
        raise DocxRenderError(f"LibreOffice did not produce expected PDF: {pdf_path}")
    return DocxRenderResult(pdf_path=pdf_path, page_count=_pdf_page_count(pdf_path))


def _rendered_pagination_metadata(
    docx_path: str | Path,
    *,
    asset_dir: Path | None,
    rendered_pdf_path: str | Path | None,
    soffice_path: str | Path | None,
    render_runner: Callable[..., Any],
) -> tuple[dict[str, Any], list[str]]:
    warnings: list[str] = []
    base = {
        "requested": "rendered",
        "effective": "logical_fallback",
        "backend": "soffice",
    }
    try:
        if rendered_pdf_path is not None:
            rendered_pdf = Path(rendered_pdf_path).resolve()
            result = DocxRenderResult(pdf_path=rendered_pdf, page_count=_pdf_page_count(rendered_pdf))
        else:
            output_dir = asset_dir or (Path(docx_path).resolve().parent / ".docx_rendered")
            result = render_docx_to_pdf(
                docx_path,
                output_dir=output_dir,
                soffice_path=soffice_path,
                runner=render_runner,
            )
    except DocxRenderError as exc:
        warning = _redact_local_paths(f"DOCX rendered pagination unavailable: {exc}")
        warnings.append(warning)
        return {**base, "reason": "render_unavailable", "warnings": warnings}, warnings

    if result.page_count <= 1:
        return {
            **base,
            "effective": "rendered_single_page",
            "rendered_pdf_path": str(result.pdf_path),
            "rendered_page_count": result.page_count,
        }, warnings
    warning = "DOCX rendered page count available, but block-to-page layout mapping is unavailable; using logical page fallback."
    warnings.append(warning)
    return {
        **base,
        "reason": "missing_layout_mapping",
        "rendered_pdf_path": str(result.pdf_path),
        "rendered_page_count": result.page_count,
        "warnings": warnings,
    }, warnings


def _extract_images(
    document,
    *,
    doc_id: str,
    asset_dir: Path | None,
    page_anchor_type: str,
    rendered_page_no: int | None,
) -> list[DocumentAsset]:
    assets: list[DocumentAsset] = []
    for rel_id, rel in document.part.rels.items():
        target_part = getattr(rel, "target_part", None)
        content_type = str(getattr(target_part, "content_type", ""))
        if not content_type.startswith("image/"):
            continue
        asset_id = make_asset_id(doc_id, 1, len(assets) + 1, "figure")
        asset_path = f"docx://{doc_id}/image/{asset_id}"
        if asset_dir is not None:
            output_path = asset_dir / f"{asset_id}{_image_extension(content_type)}"
            output_path.parent.mkdir(parents=True, exist_ok=True)
            output_path.write_bytes(target_part.blob)
            asset_path = str(output_path)
        assets.append(
            DocumentAsset(
                asset_id=asset_id,
                doc_id=doc_id,
                page_no=1,
                object_type="figure",
                asset_path=asset_path,
                text="",
                confidence=1.0,
                metadata={
                    "source": "docx_image_relationship",
                    "content_type": content_type,
                    **asset_anchor_metadata(
                        source="docx_image_relationship",
                        page_no=1,
                        page_anchor_type=page_anchor_type,
                        rendered_page_no=rendered_page_no,
                        relationship_id=str(rel_id),
                    ),
                },
            )
        )
    return assets


def extract_docx_text(path: str | Path) -> str:
    """Return DOCX paragraph/table text for classification."""
    document = _load_document(path)
    return normalize_text("\n".join(_text_parts_from_document(document)))


def inspect_docx(path: str | Path) -> tuple[str, int]:
    """Return text and logical page count for DOCX classification."""
    document = _load_document(path)
    text = normalize_text("\n".join(_text_parts_from_document(document)))
    if text:
        return text, 1
    return "", 1 if document.part.rels else 0


def parse_docx(
    docx_path: str | Path,
    *,
    doc_id: str,
    asset_dir: str | Path | None = None,
    extract_image_assets: bool = True,
    pagination_mode: str = "logical",
    rendered_pdf_path: str | Path | None = None,
    soffice_path: str | Path | None = None,
    render_runner: Callable[..., Any] = subprocess.run,
) -> list[OcrPage]:
    """Parse a DOCX into page records compatible with PDF OCR output.

    ``pagination_mode="rendered"`` asks for a headless LibreOffice page-count
    pass. Without block-level rendered layout mapping, multi-page DOCX files
    deliberately fall back to the Phase-3 logical page while recording why.
    """
    if pagination_mode not in {"logical", "rendered"}:
        raise ValueError("pagination_mode must be 'logical' or 'rendered'.")

    document = _load_document(docx_path)
    resolved_asset_dir = Path(asset_dir).resolve() if asset_dir is not None else None
    pagination: dict[str, Any] = {"requested": pagination_mode, "effective": "logical", "backend": "python-docx"}
    warnings: list[str] = []
    if pagination_mode == "rendered":
        pagination, warnings = _rendered_pagination_metadata(
            docx_path,
            asset_dir=resolved_asset_dir,
            rendered_pdf_path=rendered_pdf_path,
            soffice_path=soffice_path,
            render_runner=render_runner,
        )
    page_anchor_type = str(pagination.get("effective") or "logical")
    rendered_page_no = 1 if page_anchor_type == "rendered_single_page" else None
    blocks: list[PageBlock] = []
    body_parts: list[str] = []

    for paragraph in document.paragraphs:
        text = normalize_text(paragraph.text)
        if not text:
            continue
        block_index = len(blocks) + 1
        block_type = _paragraph_kind(text, paragraph, block_index=block_index)
        blocks.append(
            PageBlock(
                block_id=f"p0001_b{block_index:04d}",
                text=text,
                block_type=block_type,
                confidence=1.0,
                metadata={
                    "source": "docx_paragraph",
                    "style": getattr(paragraph.style, "name", None),
                    "page_anchor_type": page_anchor_type,
                },
            )
        )
        if block_type in {"body", "title"}:
            body_parts.append(text)

    table_index = 0
    for table_text in _iter_table_texts(document):
        table_index += 1
        block_index = len(blocks) + 1
        asset_id = make_asset_id(doc_id, 1, table_index, "table")
        asset_path = f"docx://{doc_id}/table/{asset_id}"
        blocks.append(
            PageBlock(
                block_id=f"p0001_b{block_index:04d}",
                text=table_text,
                block_type="table",
                confidence=1.0,
                asset_id=asset_id,
                asset_path=asset_path,
                metadata={"source": "docx_table"},
            )
        )
        body_parts.append(table_text)

    assets: list[DocumentAsset] = []
    for block in blocks:
        if block.block_type == "table" and block.asset_id:
            assets.append(
                DocumentAsset(
                    asset_id=block.asset_id,
                    doc_id=doc_id,
                    page_no=1,
                    object_type="table",
                    asset_path=block.asset_path,
                    text=block.text,
                    confidence=1.0,
                    metadata={
                        "block_id": block.block_id,
                        "source": "docx_table",
                        **asset_anchor_metadata(
                            source="docx_table",
                            page_no=1,
                            page_anchor_type=page_anchor_type,
                            block_id=block.block_id,
                            rendered_page_no=rendered_page_no,
                        ),
                    },
                )
            )

    if extract_image_assets:
        assets.extend(
            _extract_images(
                document,
                doc_id=doc_id,
                asset_dir=resolved_asset_dir,
                page_anchor_type=page_anchor_type,
                rendered_page_no=rendered_page_no,
            )
        )

    normalized_text = normalize_text("\n".join(body_parts))
    if not normalized_text and not assets:
        raise DocxParseError("DOCX contains no extractable text, tables, or images.")

    return [
        OcrPage(
            doc_id=doc_id,
            page_no=1,
            primary_engine="python-docx",
            fallback_used=False,
            ocr_confidence=1.0 if normalized_text or assets else None,
            layout_quality="ok" if normalized_text else "low",
            raw_text=normalize_text("\n".join(block.text for block in blocks)),
            normalized_text=normalized_text,
            blocks=blocks,
            assets=assets,
            needs_review=not bool(normalized_text),
            metadata={"docx_pagination": pagination, "warnings": warnings},
        )
    ]


def docx_page_count(path: str | Path) -> int:
    return inspect_docx(path)[1]
